from __future__ import annotations
import io
import os
import re
import tempfile
import urllib.parse
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

import cv2
import numpy as np
from PIL import Image, ImageOps

def is_image_file(filename: str, mime: str = "") -> bool:
    ext = Path(filename).suffix.lower()
    return ext in {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".tif", ".heic", ".heif"} or "image/" in mime

def is_pdf_file(filename: str, mime: str = "") -> bool:
    ext = Path(filename).suffix.lower()
    return ext == ".pdf" or "pdf" in mime

def is_docx_file(filename: str, mime: str = "") -> bool:
    ext = Path(filename).suffix.lower()
    return ext in {".docx", ".doc"} or "word" in mime or "officedocument" in mime

def is_video_file(filename: str, mime: str = "") -> bool:
    ext = Path(filename).suffix.lower()
    return ext in {".mp4", ".mov", ".avi", ".webm", ".mkv", ".m4v", ".flv"} or "video/" in mime

def is_audio_file(filename: str, mime: str = "") -> bool:
    ext = Path(filename).suffix.lower()
    return ext in {".wav", ".mp3", ".ogg", ".m4a", ".aac", ".flac", ".wma"} or "audio/" in mime

def parse_image_bytes(raw_bytes: bytes, filename: str = "image.png") -> Tuple[np.ndarray, Dict[str, Any]]:
    arr = np.frombuffer(raw_bytes, dtype=np.uint8)
    image = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    
    metadata = {
        "filename": filename,
        "media_type": "image",
        "file_size_bytes": len(raw_bytes),
        "dimensions": None,
        "has_exif": False,
        "exif_details": {}
    }
    
    if image is None:
        try:
            pil_img = Image.open(io.BytesIO(raw_bytes))
            pil_img = ImageOps.exif_transpose(pil_img)
            rgb_arr = np.array(pil_img.convert("RGB"))
            image = cv2.cvtColor(rgb_arr, cv2.COLOR_RGB2BGR)
        except Exception:
            raise ValueError("Unable to decode image file. Unsupported or corrupted format.")

    h, w = image.shape[:2]
    metadata["dimensions"] = {"width": w, "height": h, "channels": 3}
    
    # Try reading EXIF
    try:
        pil_img = Image.open(io.BytesIO(raw_bytes))
        exif = pil_img.getexif()
        if exif:
            metadata["has_exif"] = True
            details = {}
            for tag_id, val in exif.items():
                tag_name = str(tag_id)
                if isinstance(val, (int, float, str)):
                    details[tag_name] = str(val)[:50]
            metadata["exif_details"] = details
    except Exception:
        pass

    return image, metadata

def parse_pdf_bytes(raw_bytes: bytes, filename: str = "document.pdf") -> Tuple[List[np.ndarray], str, Dict[str, Any]]:
    images = []
    text_content = []
    metadata = {
        "filename": filename,
        "media_type": "pdf",
        "file_size_bytes": len(raw_bytes),
        "total_pages": 0,
        "extracted_images_count": 0
    }
    
    # Try pymupdf (fitz)
    try:
        import pymupdf as fitz
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        metadata["total_pages"] = len(doc)
        
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            text_content.append(page.get_text())
            
            # Render page to high-res pixmap
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            arr = np.frombuffer(img_bytes, dtype=np.uint8)
            cv_img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
            if cv_img is not None:
                images.append(cv_img)
                
            # Also extract embedded images if any
            img_list = page.get_images(full=True)
            metadata["extracted_images_count"] += len(img_list)
            for img_info in img_list:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if base_image and "image" in base_image:
                    sub_arr = np.frombuffer(base_image["image"], dtype=np.uint8)
                    sub_cv = cv2.imdecode(sub_arr, cv2.IMREAD_COLOR)
                    if sub_cv is not None and sub_cv.shape[0] > 100 and sub_cv.shape[1] > 100:
                        images.append(sub_cv)
                        
    except Exception as e:
        # Fallback using pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            metadata["total_pages"] = len(reader.pages)
            for page in reader.pages:
                text_content.append(page.extract_text() or "")
                for img_obj in page.images:
                    sub_arr = np.frombuffer(img_obj.data, dtype=np.uint8)
                    sub_cv = cv2.imdecode(sub_arr, cv2.IMREAD_COLOR)
                    if sub_cv is not None:
                        images.append(sub_cv)
        except Exception:
            pass

    full_text = "\n".join(text_content).strip()
    return images, full_text, metadata

def parse_docx_bytes(raw_bytes: bytes, filename: str = "document.docx") -> Tuple[List[np.ndarray], str, Dict[str, Any]]:
    images = []
    text_content = []
    metadata = {
        "filename": filename,
        "media_type": "docx",
        "file_size_bytes": len(raw_bytes),
        "paragraphs_count": 0,
        "extracted_images_count": 0
    }
    
    try:
        import docx
        doc = docx.Document(io.BytesIO(raw_bytes))
        for p in doc.paragraphs:
            if p.text:
                text_content.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text:
                        text_content.append(cell.text)
        metadata["paragraphs_count"] = len(text_content)
        
        # Extract inline images from docx package
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref.lower():
                try:
                    img_data = rel.target_part.blob
                    arr = np.frombuffer(img_data, dtype=np.uint8)
                    cv_img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
                    if cv_img is not None and cv_img.shape[0] > 80 and cv_img.shape[1] > 80:
                        images.append(cv_img)
                except Exception:
                    continue
        metadata["extracted_images_count"] = len(images)
    except Exception as e:
        text_content.append(f"Docx parsing error: {e}")

    full_text = "\n".join(text_content).strip()
    return images, full_text, metadata

def parse_video_bytes(raw_bytes: bytes, filename: str = "video.mp4") -> Tuple[List[np.ndarray], Dict[str, Any]]:
    metadata = {
        "filename": filename,
        "media_type": "video",
        "file_size_bytes": len(raw_bytes),
        "fps": 0.0,
        "total_frames": 0,
        "duration_seconds": 0.0,
        "video_category": "Standard Video",
        "temporal_consistency": 1.0,
        "motion_variance": 0.0,
        "deepfake_flicker_score": 0.0,
        "quality_signals": []
    }
    
    frames = []
    suffix = Path(filename).suffix or ".mp4"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tf:
        tf.write(raw_bytes)
        tf_path = tf.name

    try:
        cap = cv2.VideoCapture(tf_path)
        fps = cap.get(cv2.CAP_PROP_FPS) or 30.0
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        metadata["fps"] = round(fps, 2)
        metadata["total_frames"] = total_frames
        metadata["duration_seconds"] = round(total_frames / max(1.0, fps), 2) if total_frames > 0 else 0.0

        # Sample up to 16 keyframes
        sample_step = max(1, total_frames // 16) if total_frames > 16 else 1
        frame_idx = 0
        diff_scores = []
        last_gray = None

        while True:
            ret, frame = cap.read()
            if not ret or frame is None:
                break
            
            if frame_idx % sample_step == 0 and len(frames) < 16:
                frames.append(frame.copy())
            
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            if last_gray is not None:
                # Compute absolute diff between consecutive frames
                diff = cv2.absdiff(gray, last_gray)
                diff_scores.append(float(diff.mean()))
            last_gray = gray
            frame_idx += 1
            if frame_idx > 300: # limit to first 300 frames for speed
                break
                
        cap.release()

        # Categorize: Timelapse vs Slow Motion vs Normal
        if diff_scores:
            mean_diff = float(np.mean(diff_scores))
            var_diff = float(np.var(diff_scores))
            metadata["motion_variance"] = round(var_diff, 3)

            if fps > 50 or (mean_diff < 1.2 and metadata["duration_seconds"] > 8):
                metadata["video_category"] = "Slow Motion Video"
                metadata["quality_signals"].append("slow_motion_framerate")
            elif fps < 15 or mean_diff > 35.0:
                metadata["video_category"] = "Timelapse / High Velocity Video"
                metadata["quality_signals"].append("timelapse_velocity")
            else:
                metadata["video_category"] = "Standard Video"

            # Flicker / Temporal anomaly check (deepfakes often exhibit high inter-frame flickering)
            flicker = min(1.0, var_diff / 100.0) if var_diff > 30 else 0.05
            metadata["deepfake_flicker_score"] = round(flicker, 3)
            metadata["temporal_consistency"] = round(max(0.0, 1.0 - (flicker * 0.7)), 3)
            if flicker > 0.4:
                metadata["quality_signals"].append("high_temporal_flicker")
                
    finally:
        if os.path.exists(tf_path):
            try:
                os.remove(tf_path)
            except Exception:
                pass

    return frames, metadata

def parse_audio_bytes(raw_bytes: bytes, filename: str = "audio.wav") -> Dict[str, Any]:
    metadata = {
        "filename": filename,
        "media_type": "audio",
        "file_size_bytes": len(raw_bytes),
        "duration_seconds": 0.0,
        "sample_rate": 0,
        "channels": 1,
        "synthetic_voice_probability": 0.05,
        "spectral_flatness": 0.0,
        "zero_crossing_rate": 0.0,
        "audio_signals": [],
        "verdict": "GENUINE_HUMAN_VOICE"
    }

    try:
        # Try wave module first
        import wave
        with wave.open(io.BytesIO(raw_bytes), "rb") as wf:
            n_channels = wf.getnchannels()
            sampwidth = wf.getsampwidth()
            framerate = wf.getframerate()
            n_frames = wf.getnframes()
            frames_raw = wf.readframes(n_frames)
            
            metadata["sample_rate"] = framerate
            metadata["channels"] = n_channels
            metadata["duration_seconds"] = round(n_frames / max(1.0, framerate), 2)
            
            # Convert to numpy array
            if sampwidth == 2:
                audio_data = np.frombuffer(frames_raw, dtype=np.int16).astype(np.float32)
            else:
                audio_data = np.frombuffer(frames_raw, dtype=np.uint8).astype(np.float32) - 128
                
            if n_channels > 1:
                audio_data = audio_data[::n_channels] # mono channel
                
            # Audio signal metrics
            if len(audio_data) > 100:
                # Normalization
                norm_data = audio_data / (np.max(np.abs(audio_data)) + 1e-6)
                
                # Zero Crossing Rate (ZCR)
                zcr = float(np.mean(np.abs(np.diff(np.sign(norm_data)))) / 2.0)
                metadata["zero_crossing_rate"] = round(zcr, 4)
                
                # FFT frequency distribution
                fft_vals = np.abs(np.fft.rfft(norm_data[:min(len(norm_data), 32768)]))
                fft_norm = fft_vals / (np.sum(fft_vals) + 1e-8)
                
                # Spectral flatness (ratio of geometric mean to arithmetic mean)
                geom_mean = np.exp(np.mean(np.log(fft_norm + 1e-12)))
                arith_mean = np.mean(fft_norm)
                flatness = float(geom_mean / max(arith_mean, 1e-8))
                metadata["spectral_flatness"] = round(flatness, 4)
                
                # Vocal Intonation / Pitch Jitter Analysis (Auto-correlation across 50ms windows)
                win_size = int(framerate * 0.05)
                lags = []
                for i in range(0, len(norm_data) - win_size, max(1, win_size // 2)):
                    w = norm_data[i:i+win_size]
                    corr = np.correlate(w, w, mode='full')[win_size-1:]
                    min_lag = max(20, int(framerate / 400)) # 400Hz ceiling
                    max_lag = min(len(corr), int(framerate / 50))  # 50Hz floor
                    if max_lag > min_lag:
                        peak_lag = min_lag + np.argmax(corr[min_lag:max_lag])
                        lags.append(peak_lag)
                
                pitch_std = float(np.std(lags)) if len(lags) > 3 else 0.0
                metadata["pitch_variance"] = round(pitch_std, 2)

                # Synthetic speech detectors: TTS / Vocoder artifacts
                synth_score = 0.05
                if pitch_std < 1.2:
                    synth_score += 0.75
                    metadata["audio_signals"].append("rigid_unnatural_pitch_profile")
                if flatness < 0.0001:
                    synth_score += 0.40
                    metadata["audio_signals"].append("robotic_spectral_peaks")
                if zcr < 0.01 or zcr > 0.35:
                    synth_score += 0.30
                    metadata["audio_signals"].append("atypical_zero_crossing_rate")
                if metadata["duration_seconds"] < 0.5:
                    metadata["audio_signals"].append("insufficient_audio_length")

                metadata["synthetic_voice_probability"] = round(min(0.99, synth_score), 2)
                if metadata["synthetic_voice_probability"] >= 0.40:
                    metadata["verdict"] = "SUSPICIOUS_AI_GENERATED_VOICE"
                else:
                    metadata["verdict"] = "GENUINE_HUMAN_VOICE"
    except Exception as e:
        metadata["audio_signals"].append(f"audio_codec_fallback: {str(e)[:30]}")
        metadata["duration_seconds"] = round(len(raw_bytes) / 32000.0, 2)
        metadata["synthetic_voice_probability"] = 0.12

    return metadata

def fetch_url_resource(url: str, timeout: int = 10) -> Tuple[bytes, str, str]:
    import requests
    parsed = urllib.parse.urlparse(url)
    if not parsed.scheme or not parsed.netloc:
        raise ValueError("Invalid URL scheme. Must start with http:// or https://")
        
    headers = {
        "User-Agent": "Mozilla/5.0 (The-Imposter-Check/2.0 Identity-Verification-Bot)"
    }
    resp = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True)
    resp.raise_for_status()
    
    raw_bytes = resp.content
    content_type = resp.headers.get("Content-Type", "").lower()
    
    # Infer filename from URL path
    path = parsed.path.strip("/")
    filename = Path(path).name if path else "downloaded_asset"
    if not Path(filename).suffix:
        if "jpeg" in content_type or "jpg" in content_type: filename += ".jpg"
        elif "png" in content_type: filename += ".png"
        elif "pdf" in content_type: filename += ".pdf"
        elif "docx" in content_type: filename += ".docx"
        elif "mp4" in content_type: filename += ".mp4"
        elif "html" in content_type: filename += ".html"
        else: filename += ".bin"

    return raw_bytes, filename, content_type
