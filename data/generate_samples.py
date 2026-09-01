"""
Generator script to create comprehensive sample test files for:
- Genuine & Tampered Photos
- Multi-Card images (2 IDs in 1 file)
- PDF identity dossiers
- DOCX verification documents
- Genuine & AI-synthesized Audio WAVs
- Video test clips (standard, timelapse, slowmo)
- Cross-comparison test sets (Matching vs Imposter)
"""

import os
import math
import wave
import struct
from pathlib import Path
import cv2
import numpy as np
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "data" / "sample_inputs"
OUT_DIR.mkdir(parents=True, exist_ok=True)

print(f"Generating sample files in: {OUT_DIR}")

# -------------------------------------------------------------
# 1. HELPER: DRAW SYNTHETIC FACE
# -------------------------------------------------------------
def draw_synthetic_face(w=120, h=140, is_alternate=False):
    canvas = np.full((h, w, 3), (225, 230, 235), dtype=np.uint8)
    
    # Skin tone
    skin_color = (180, 205, 235) if not is_alternate else (160, 190, 220)
    center = (w // 2, h // 2 + 5)
    
    # Face shape
    cv2.ellipse(canvas, center, (w // 3, h // 3 + 5), 0, 0, 360, skin_color, -1)
    
    # Hair
    hair_color = (40, 30, 25) if not is_alternate else (30, 60, 110)
    cv2.ellipse(canvas, (w // 2, h // 3), (w // 3 + 2, h // 4), 0, 180, 360, hair_color, -1)
    
    # Eyes
    eye_y = center[1] - 10
    left_eye = (center[0] - 16, eye_y)
    right_eye = (center[0] + 16, eye_y)
    
    cv2.circle(canvas, left_eye, 5, (255, 255, 255), -1)
    cv2.circle(canvas, right_eye, 5, (255, 255, 255), -1)
    cv2.circle(canvas, left_eye, 2, (30, 30, 30), -1)
    cv2.circle(canvas, right_eye, 2, (30, 30, 30), -1)
    
    # Eyebrows
    cv2.line(canvas, (left_eye[0] - 8, eye_y - 8), (left_eye[0] + 8, eye_y - 7), hair_color, 2)
    cv2.line(canvas, (right_eye[0] - 8, eye_y - 7), (right_eye[0] + 8, eye_y - 8), hair_color, 2)
    
    # Nose
    cv2.line(canvas, (center[0], center[1] - 4), (center[0] - 2, center[1] + 8), (140, 170, 200), 2)
    cv2.line(canvas, (center[0] - 2, center[1] + 8), (center[0] + 3, center[1] + 8), (140, 170, 200), 2)
    
    # Mouth
    cv2.ellipse(canvas, (center[0], center[1] + 20), (12, 5), 0, 0, 180, (90, 100, 190), 2)
    
    # Body collar
    cv2.ellipse(canvas, (w // 2, h + 15), (w // 2, 25), 0, 180, 360, (120, 60, 40), -1)
    return canvas

# -------------------------------------------------------------
# 2. IMAGE GENERATORS: GENUINE, TAMPERED & MULTI-CARD
# -------------------------------------------------------------
def make_id_card(name="MAHITA", doc_id="DL-8892140A", title="FEDERAL DRIVER LICENSE", is_tampered=False, face_alternate=False):
    # Standard ID card ratio ~ 1.58 (600x380)
    w, h = 600, 380
    card = np.full((h, w, 3), (245, 248, 250), dtype=np.uint8)
    
    # Header bar
    header_color = (180, 80, 20) if not is_tampered else (60, 60, 180)
    cv2.rectangle(card, (0, 0), (w, 60), header_color, -1)
    cv2.putText(card, title, (20, 40), cv2.FONT_HERSHEY_DUPLEX, 0.8, (255, 255, 255), 2)
    
    # Golden Security Hologram Chip
    cv2.rectangle(card, (30, 85), (85, 130), (80, 190, 240), -1)
    cv2.rectangle(card, (30, 85), (85, 130), (50, 140, 180), 2)
    cv2.line(card, (30, 107), (85, 107), (50, 140, 180), 1)
    cv2.line(card, (57, 85), (57, 130), (50, 140, 180), 1)
    
    # Add Face Photo
    face = draw_synthetic_face(110, 135, is_alternate=face_alternate)
    card[150:285, 30:140] = face
    cv2.rectangle(card, (28, 148), (142, 287), (150, 150, 150), 2)
    
    # Text Fields
    cv2.putText(card, f"NAME: {name}", (170, 105), cv2.FONT_HERSHEY_DUPLEX, 0.65, (30, 30, 30), 2)
    cv2.putText(card, f"ID NO: {doc_id}", (170, 140), cv2.FONT_HERSHEY_DUPLEX, 0.6, (40, 40, 40), 1)
    cv2.putText(card, "DOB: 12 MAY 1994", (170, 175), cv2.FONT_HERSHEY_DUPLEX, 0.55, (50, 50, 50), 1)
    cv2.putText(card, "EXPIRY: 30 DEC 2032", (170, 210), cv2.FONT_HERSHEY_DUPLEX, 0.55, (50, 50, 50), 1)
    cv2.putText(card, "STATUS: CITIZEN VERIFIED", (170, 245), cv2.FONT_HERSHEY_DUPLEX, 0.55, (30, 120, 40), 2)
    
    # Machine Readable Zone (MRZ)
    cv2.rectangle(card, (0, 310), (w, h), (230, 235, 240), -1)
    cv2.putText(card, f"I<USA{doc_id}<<<<<<<<<<<<<<<", (20, 335), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (60, 60, 60), 1)
    cv2.putText(card, f"9405125M3212308USA<<<<<<<<<<<4", (20, 360), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (60, 60, 60), 1)
    
    # Tampering injection: digitally spliced patch with different noise & compression
    if is_tampered:
        # Patch a forged name on top with high noise
        patch = np.full((32, 220, 3), (255, 255, 200), dtype=np.uint8)
        noise = np.random.randint(-30, 30, patch.shape, dtype=np.int16)
        patch = np.clip(patch.astype(np.int16) + noise, 0, 255).astype(np.uint8)
        cv2.putText(patch, "FORGED: J. IMPOSTER", (5, 22), cv2.FONT_HERSHEY_DUPLEX, 0.55, (180, 0, 0), 2)
        card[88:120, 165:385] = patch
        cv2.rectangle(card, (165, 88), (385, 120), (0, 0, 255), 1)
        
    return card

# 1. Genuine ID Photo
img_genuine = make_id_card("MAHITA", "DL-8892140A", "STATE OF CALIFORNIA DRIVER LICENSE", is_tampered=False)
cv2.imwrite(str(OUT_DIR / "1_genuine_id_card.png"), img_genuine)

# 2. Forged Tampered ID
img_tampered = make_id_card("MAHITA", "DL-8892140A", "STATE OF CALIFORNIA DRIVER LICENSE", is_tampered=True)
cv2.imwrite(str(OUT_DIR / "2_forged_tampered_id.png"), img_tampered)

# 3. Multi-Card Single File (Passport + Driver License in 1 image)
multi_canvas = np.full((760, 960, 3), (40, 45, 50), dtype=np.uint8)
# Add desk background texture
grid = np.zeros((760, 960, 3), dtype=np.uint8)
grid[::40, :, :] = (60, 65, 70)
grid[:, ::40, :] = (60, 65, 70)
multi_canvas = cv2.addWeighted(multi_canvas, 0.8, grid, 0.2, 0)

# Place Card #1 (Passport Card on top)
c1 = make_id_card("MAHITA", "USA-PASS-449102", "UNITED STATES PASSPORT CARD", face_alternate=False)
c1_scaled = cv2.resize(c1, (440, 280))
multi_canvas[60:340, 60:500] = c1_scaled

# Place Card #2 (Driver License on bottom right)
c2 = make_id_card("MAHITA", "DL-8892140A", "NATIONAL IDENTIFICATION CARD", face_alternate=False)
c2_scaled = cv2.resize(c2, (440, 280))
multi_canvas[400:680, 460:900] = c2_scaled

cv2.putText(multi_canvas, "CONFIDENTIAL IDENTITY VERIFICATION DESK", (60, 40), cv2.FONT_HERSHEY_DUPLEX, 0.6, (180, 180, 180), 1)
cv2.imwrite(str(OUT_DIR / "3_multi_card_single_file.png"), multi_canvas)

# -------------------------------------------------------------
# 3. PDF GENERATOR: IDENTITY DOSSIER
# -------------------------------------------------------------
try:
    import pymupdf as fitz
    pdf_doc = fitz.open()
    
    # Page 1: Identity Profile
    page1 = pdf_doc.new_page(width=595, height=842) # A4
    page1.insert_text((50, 60), "OFFICIAL IDENTITY VERIFICATION DOSSIER", fontsize=16, color=(0.1, 0.2, 0.4))
    page1.insert_text((50, 90), "Subject: MAHITA  |  Verification Status: KYC CLEARED", fontsize=11, color=(0.2, 0.5, 0.2))
    page1.insert_text((50, 120), "Date of Birth: 1994-05-12   |   Nationality: USA   |   Document Ref: DL-8892140A", fontsize=10)
    page1.insert_text((50, 145), "Biometric verification performed across facial scans and optical character recognition.", fontsize=10)
    
    # Insert card image into page 1
    _, buf = cv2.imencode('.png', img_genuine)
    page1.insert_image(fitz.Rect(50, 180, 545, 490), stream=buf.tobytes())
    page1.insert_text((50, 520), "Security Features Inspected: Microprint, Optical Chip, Holographic Overlay.", fontsize=9, color=(0.4, 0.4, 0.4))
    
    # Page 2: Secondary Proof of Address
    page2 = pdf_doc.new_page(width=595, height=842)
    page2.insert_text((50, 60), "UTILITY BILL & RESIDENTIAL CONFIRMATION", fontsize=14, color=(0.1, 0.2, 0.4))
    page2.insert_text((50, 100), "Resident Name: MAHITA\nAddress: 742 Evergreen Terrace, Springfield, OR 97477\nBilling Period: August 2026\nAccount Number: UTIL-992140-A\nStatus: Paid in Full", fontsize=11)
    
    pdf_path = OUT_DIR / "4_identity_dossier.pdf"
    pdf_doc.save(str(pdf_path))
    pdf_doc.close()
    print(f"Created PDF: {pdf_path}")
except Exception as e:
    print(f"PDF creation error: {e}")

# -------------------------------------------------------------
# 4. DOCX GENERATOR: EMPLOYMENT IDENTITY FORM
# -------------------------------------------------------------
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    
    doc = docx.Document()
    title = doc.add_heading("Corporate Identity & Employment Verification", 0)
    
    p = doc.add_paragraph()
    p.add_run("Verified Employee: ").bold = True
    p.add_run("Mahita\n")
    p.add_run("Department: ").bold = True
    p.add_run("Cybersecurity Threat Intelligence\n")
    p.add_run("Employee ID: ").bold = True
    p.add_run("EMP-889214-INT\n")
    p.add_run("Clearance Level: ").bold = True
    p.add_run("Top Secret / SCI Level 4")
    
    # Add Table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Document Type"
    table.rows[0].cells[1].text = "Identification Number"
    table.rows[1].cells[0].text = "Driver License"
    table.rows[1].cells[1].text = "DL-8892140A"
    table.rows[2].cells[0].text = "Passport Card"
    table.rows[2].cells[1].text = "USA-PASS-449102"
    
    # Save temp card image and insert into docx
    temp_img_path = OUT_DIR / "temp_card_for_docx.png"
    cv2.imwrite(str(temp_img_path), img_genuine)
    doc.add_paragraph("\nAttached Government Credential:")
    doc.add_picture(str(temp_img_path), width=Inches(4.5))
    if temp_img_path.exists():
        temp_img_path.unlink()
        
    docx_path = OUT_DIR / "5_employment_identity.docx"
    doc.save(str(docx_path))
    print(f"Created DOCX: {docx_path}")
except Exception as e:
    print(f"DOCX creation error: {e}")

# -------------------------------------------------------------
# 5. AUDIO GENERATOR: GENUINE SPEECH VS SYNTHETIC AI DEEPFAKE
# -------------------------------------------------------------
def generate_audio_wav(filename, is_synthetic=False, duration=3.0, sample_rate=16000):
    n_samples = int(duration * sample_rate)
    samples = []
    
    base_freq = 150.0 # Fundamental pitch ~150Hz
    
    for i in range(n_samples):
        t = i / sample_rate
        if not is_synthetic:
            # Natural voice simulation: subtle pitch jitter, formant harmonics, envelope
            pitch = base_freq + 12.0 * math.sin(2 * math.pi * 3.5 * t) + 4.0 * math.sin(2 * math.pi * 7.0 * t)
            harm1 = 0.6 * math.sin(2 * math.pi * pitch * t)
            harm2 = 0.3 * math.sin(2 * math.pi * (pitch * 2.0) * t)
            harm3 = 0.15 * math.sin(2 * math.pi * (pitch * 3.0) * t)
            formant = math.sin(2 * math.pi * 800 * t) * 0.1 * math.sin(2 * math.pi * 4 * t)
            sample_val = (harm1 + harm2 + harm3 + formant) * min(1.0, min(t * 5, (duration - t) * 5))
        else:
            # Synthetic / Robotic TTS artifact: constant pitch, flat spectral spikes, harmonic buzz
            pitch = 180.0 # perfectly flat pitch without human vocal micro-vibrato
            buzz = math.sin(2 * math.pi * pitch * t) + 0.8 * math.sin(2 * math.pi * (pitch * 2) * t) + 0.6 * math.sin(2 * math.pi * (pitch * 4) * t)
            sample_val = buzz * 0.5
            
        int_sample = int(np.clip(sample_val * 24000, -32000, 32000))
        samples.append(int_sample)
        
    wav_path = OUT_DIR / filename
    with wave.open(str(wav_path), "wb") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sample_rate)
        data = struct.pack(f"<{len(samples)}h", *samples)
        wf.writeframes(data)
    print(f"Created Audio: {wav_path}")

generate_audio_wav("6_genuine_voice_sample.wav", is_synthetic=False)
generate_audio_wav("7_ai_deepfake_voice_sample.wav", is_synthetic=True)

# -------------------------------------------------------------
# 6. VIDEO GENERATORS: LIVENESS SELFIE, TIMELAPSE & SLOWMO
# -------------------------------------------------------------
def generate_test_video(filename, mode="liveness", fps=30.0, n_frames=90):
    w, h = 640, 480
    video_path = OUT_DIR / filename
    fourcc = cv2.VideoWriter_fourcc(*'mp4v')
    out = cv2.VideoWriter(str(video_path), fourcc, fps, (w, h))
    
    for i in range(n_frames):
        frame = np.full((h, w, 3), (25, 30, 40), dtype=np.uint8)
        
        # Draw background grid
        cv2.line(frame, (0, 240), (640, 240), (40, 45, 55), 1)
        cv2.line(frame, (320, 0), (320, 480), (40, 45, 55), 1)
        
        if mode == "liveness":
            # Moving face nodding & blinking
            dx = int(25 * math.sin(i * 0.08))
            dy = int(10 * math.cos(i * 0.06))
            face = draw_synthetic_face(200, 250, is_alternate=False)
            frame[110 + dy:360 + dy, 220 + dx:420 + dx] = face
            cv2.putText(frame, "LIVE BIOMETRIC MOTION CHECK", (30, 50), cv2.FONT_HERSHEY_DUPLEX, 0.7, (0, 242, 254), 2)
            cv2.putText(frame, f"FRAME: {i+1:03d}/090", (30, 440), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (150, 150, 150), 1)
            
        elif mode == "timelapse":
            # High speed rotating pattern & rapid color shifts
            angle = (i * 24) % 360
            center = (w // 2, h // 2)
            cv2.circle(frame, center, 140, (0, 200, 255), 3)
            rad = math.radians(angle)
            end_x = int(center[0] + 120 * math.cos(rad))
            end_y = int(center[1] + 120 * math.sin(rad))
            cv2.line(frame, center, (end_x, end_y), (0, 240, 100), 4)
            cv2.putText(frame, "TIMELAPSE VELOCITY TEST (ACCELERATED)", (30, 50), cv2.FONT_HERSHEY_DUPLEX, 0.7, (245, 158, 11), 2)
            
        elif mode == "slowmo":
            # Very subtle micro movements at high 60fps
            shift = int(5 * math.sin(i * 0.02))
            face = draw_synthetic_face(200, 250, is_alternate=False)
            frame[115:365, 220 + shift:420 + shift] = face
            cv2.putText(frame, "SLOW MOTION DYNAMICS (60 FPS CAPTURE)", (30, 50), cv2.FONT_HERSHEY_DUPLEX, 0.7, (168, 85, 247), 2)
            
        out.write(frame)
        
    out.release()
    print(f"Created Video: {video_path}")

generate_test_video("8_liveness_selfie_video.mp4", mode="liveness", fps=30.0, n_frames=90)
generate_test_video("9_timelapse_test_video.mp4", mode="timelapse", fps=10.0, n_frames=60)
generate_test_video("10_slowmo_test_video.mp4", mode="slowmo", fps=60.0, n_frames=120)

# -------------------------------------------------------------
# 7. MULTI-FILE CROSS-IDENTITY COMPARISON TEST SETS
# -------------------------------------------------------------
# Matching Set: Mahita ID + Mahita Selfie
selfie_match = np.full((500, 500, 3), (30, 35, 45), dtype=np.uint8)
face_match = draw_synthetic_face(260, 320, is_alternate=False)
selfie_match[90:410, 120:380] = face_match
cv2.putText(selfie_match, "SELFIE: MAHITA (GENUINE MATCH)", (30, 460), cv2.FONT_HERSHEY_SIMPLEX, 0.55, (16, 185, 129), 2)
cv2.imwrite(str(OUT_DIR / "11_match_mahita_selfie.png"), selfie_match)


# Imposter / Fake Set: Imposter Selfie with completely DIFFERENT Face
selfie_imposter = np.full((500, 500, 3), (30, 35, 45), dtype=np.uint8)
face_imposter = draw_synthetic_face(260, 320, is_alternate=True) # Alternate face features
selfie_imposter[90:410, 120:380] = face_imposter
cv2.putText(selfie_imposter, "IMPOSTER: DIFFERENT FACE SUBJECT", (30, 460), cv2.FONT_HERSHEY_SIMPLEX, 0.55, (239, 68, 68), 2)
cv2.imwrite(str(OUT_DIR / "12_imposter_mismatch_selfie.png"), selfie_imposter)

print(f"\nSuccessfully generated all sample input files in {OUT_DIR}!")
